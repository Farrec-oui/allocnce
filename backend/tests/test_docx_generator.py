"""Tests unitaires et d'intégration pour generate_docx."""
import os
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from services.docx_generator import generate_docx, COLOR_MAP, SEPARATOR


# ---------------------------------------------------------------------------
# Fixture : alloc_data minimal synthétique
# ---------------------------------------------------------------------------

def _flight(flt_no, dep, arr, std="0600", sta="0700",
            ac_reg="OE-IJR", ac_type="320", pax=150,
            crew_str="2+4/0", captain=None,
            crew_change_before=False, date="27.06"):
    return {
        "date": date, "flt_no": flt_no, "dep": dep, "arr": arr,
        "std": std, "sta": sta, "ac_reg": ac_reg, "ac_type": ac_type,
        "capacity": 186, "pax": pax, "crew_str": crew_str,
        "captain": captain, "crew_change_before": crew_change_before,
    }


MINIMAL_ALLOC = {
    "date": "27 JUN",
    "based": [
        {
            "ac_reg": "OE-IJR",
            "new_based": False,
            "flights": [
                _flight("EJU1687", "NCE", "OLB", "0400", "0500",
                        captain="DURENDEAU DAVID"),
                _flight("EJU1688", "OLB", "NCE", "0545", "0645",
                        captain="DURENDEAU DAVID"),
            ],
        },
        {
            "ac_reg": "OE-ICI",
            "new_based": True,
            "flights": [
                _flight("EJU1663", "NCE", "HRG", "0500", "0915",
                        captain="GIANNORSI JEAN-ROCH"),
                _flight("EJU1664", "HRG", "NCE", "1005", "1445",
                        captain="GIANNORSI JEAN-ROCH"),
                _flight("EJU1745", "NCE", "FNC", "1600", "1945",
                        captain="VIGNONI MANUEL",
                        crew_change_before=True),
            ],
        },
    ],
    "autres_rotations": [
        {
            "ac_reg": "G-UZHF",
            "new_based": False,
            "flights": [
                _flight("EZY8417", "LGW", "NCE", "0500", "0705",
                        ac_reg="G-UZHF", ac_type="320", pax=175),
                _flight("EZY8418", "NCE", "LGW", "0740", "0945",
                        ac_reg="G-UZHF", ac_type="320", pax=188),
            ],
        },
    ],
    "swap_refs": [],
    "ferry": [],
}


@pytest.fixture
def docx_path(tmp_path):
    out = str(tmp_path / "test_alloc.docx")
    generate_docx(MINIMAL_ALLOC, [], out)
    return out


@pytest.fixture
def doc(docx_path):
    return Document(docx_path)


@pytest.fixture
def paras(doc):
    return doc.paragraphs


# ---------------------------------------------------------------------------
# Création du fichier
# ---------------------------------------------------------------------------

def test_fichier_cree(docx_path):
    assert os.path.exists(docx_path)
    assert os.path.getsize(docx_path) > 0


def test_retourne_chemin(tmp_path):
    out = str(tmp_path / "sub" / "alloc.docx")
    result = generate_docx(MINIMAL_ALLOC, [], out)
    assert result == out
    assert os.path.exists(out)


def test_cree_repertoires(tmp_path):
    out = str(tmp_path / "deep" / "nested" / "alloc.docx")
    generate_docx(MINIMAL_ALLOC, [], out)
    assert os.path.exists(out)


# ---------------------------------------------------------------------------
# Structure globale
# ---------------------------------------------------------------------------

def test_nb_paragraphes_non_nul(paras):
    assert len(paras) > 0


def test_titre_premier_paragraphe(paras):
    assert paras[0].text == "Operations EasyJet NCE  27 JUN Toutes les heures sont en Z"


def test_titre_police_11pt(paras):
    run = paras[0].runs[0]
    assert run.font.name == "Calibri"
    assert run.font.size == Pt(11)


def test_legende_deux_runs(paras):
    # Le deuxième paragraphe a deux runs
    assert len(paras[1].runs) == 2


def test_legende_run2_surligne_yellow(paras):
    r2 = paras[1].runs[1]
    assert r2.font.highlight_color == WD_COLOR_INDEX.YELLOW


def test_lignes_header_presentes(paras):
    texts = [p.text for p in paras]
    assert any("Night Stop J" in t for t in texts)
    assert any("Départs Remote stand" in t for t in texts)
    assert any("First Wave J+2" in t for t in texts)
    assert any("Rotations de la journée" in t for t in texts)


def test_separateur_present(paras):
    texts = [p.text for p in paras]
    assert any(t.startswith("****") for t in texts)


def test_autres_rotations_present(paras):
    texts = [p.text for p in paras]
    assert any("AUTRES ROTATIONS" in t for t in texts)


# ---------------------------------------------------------------------------
# Section Based
# ---------------------------------------------------------------------------

def test_oe_ijr_label_present(paras):
    texts = [p.text for p in paras]
    assert "OEIJR" in texts


def test_oe_ici_new_based(paras):
    texts = [p.text for p in paras]
    assert any("OEICI" in t and "New based" in t for t in texts)


def test_oe_ijr_new_based_false(paras):
    texts = [p.text for p in paras]
    assert not any("OEIJR" in t and "New based" in t for t in texts)


def test_vol_eju1687_present(paras):
    texts = [p.text for p in paras]
    assert any("EJU1687" in t for t in texts)


def test_format_vol_correct(paras):
    line = next(p.text for p in paras if "EJU1687" in p.text)
    # date flt dep arr std sta ac_reg ac_type capacity pax crew captain
    assert line == "27.06 EJU1687 NCE OLB 0400 0500 OE-IJR 320 186 150 2+4/0 DURENDEA"


def test_heure_sans_deux_points(paras):
    line = next(p.text for p in paras if "EJU1687" in p.text)
    # Les heures ne doivent pas contenir ":"
    parts = line.split()
    assert parts[4] == "0400"   # std
    assert parts[5] == "0500"   # sta


def test_captain_tronque_8_chars(paras):
    line = next(p.text for p in paras if "EJU1687" in p.text)
    # "DURENDEAU DAVID"[:8] = "DURENDEA"
    assert line.endswith("DURENDEA")


def test_captain_absent_pas_de_suffix(paras):
    alloc = {**MINIMAL_ALLOC, "based": [{
        "ac_reg": "OE-IJR", "new_based": False,
        "flights": [_flight("EJU9999", "NCE", "OLB", captain=None)],
    }], "autres_rotations": [], "swap_refs": [], "ferry": []}
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        out = f.name
    try:
        generate_docx(alloc, [], out)
        doc2 = Document(out)
        line = next(p.text for p in doc2.paragraphs if "EJU9999" in p.text)
        # Doit se terminer par crew_str, sans nom
        assert line.endswith("2+4/0")
    finally:
        os.unlink(out)


def test_crew_change_ligne_avant_vol(paras):
    texts = [p.text for p in paras]
    idx_cc = next(i for i, t in enumerate(texts) if t == "Crew change")
    # La ligne suivante doit être un vol
    assert "EJU1745" in texts[idx_cc + 1]


def test_police_corps_9pt(paras):
    # Un vol quelconque doit être en Calibri 9pt
    vol_para = next(p for p in paras if "EJU1688" in p.text)
    run = vol_para.runs[0]
    assert run.font.name == "Calibri"
    assert run.font.size == Pt(9)


# ---------------------------------------------------------------------------
# Marges
# ---------------------------------------------------------------------------

def test_marges(doc):
    from docx.shared import Cm
    # python-docx accumule un arrondi EMU de quelques unités — on tolère ±500 EMU (~0.04mm)
    sec = doc.sections[0]
    assert abs(sec.top_margin    - Cm(1)) < 500
    assert abs(sec.bottom_margin - Cm(1)) < 500
    assert abs(sec.left_margin   - Cm(2)) < 500
    assert abs(sec.right_margin  - Cm(2)) < 500


# ---------------------------------------------------------------------------
# Highlights
# ---------------------------------------------------------------------------

class TestHighlights:
    def _gen(self, tmp_path, highlights):
        out = str(tmp_path / "hl.docx")
        generate_docx(MINIMAL_ALLOC, highlights, out)
        return [p for p in Document(out).paragraphs if p.runs]

    def test_yellow(self, tmp_path):
        paras = self._gen(tmp_path, [{"flt_no": "EJU1687", "color_index": 0}])
        p = next(p for p in paras if "EJU1687" in p.text)
        assert p.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW

    def test_bright_green(self, tmp_path):
        paras = self._gen(tmp_path, [{"flt_no": "EJU1663", "color_index": 1}])
        p = next(p for p in paras if "EJU1663" in p.text)
        assert p.runs[0].font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN

    def test_turquoise(self, tmp_path):
        paras = self._gen(tmp_path, [{"flt_no": "EZY8417", "color_index": 2}])
        p = next(p for p in paras if "EZY8417" in p.text)
        assert p.runs[0].font.highlight_color == WD_COLOR_INDEX.TURQUOISE

    def test_pink(self, tmp_path):
        paras = self._gen(tmp_path, [{"flt_no": "EZY8418", "color_index": 3}])
        p = next(p for p in paras if "EZY8418" in p.text)
        assert p.runs[0].font.highlight_color == WD_COLOR_INDEX.PINK

    def test_vol_non_surligne(self, tmp_path):
        paras = self._gen(tmp_path, [{"flt_no": "EJU1687", "color_index": 0}])
        p = next(p for p in paras if "EJU1688" in p.text)
        # EJU1688 n'est pas dans highlights → pas de surbrillance
        assert p.runs[0].font.highlight_color != WD_COLOR_INDEX.YELLOW

    def test_highlights_vide(self, tmp_path):
        paras = self._gen(tmp_path, [])
        vol = next(p for p in paras if "EJU1687" in p.text)
        assert vol.runs[0].font.highlight_color not in (
            WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.BRIGHT_GREEN,
            WD_COLOR_INDEX.TURQUOISE, WD_COLOR_INDEX.PINK,
        )

    def test_color_index_modulo(self, tmp_path):
        # index 4 → même couleur que 0 (YELLOW)
        paras = self._gen(tmp_path, [{"flt_no": "EJU1687", "color_index": 4}])
        p = next(p for p in paras if "EJU1687" in p.text)
        assert p.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW


# ---------------------------------------------------------------------------
# Section SWAP
# ---------------------------------------------------------------------------

def test_swap_nil_si_aucun(tmp_path):
    """Sans avion en swap → 'SWAP: NIL' doit apparaître."""
    out = str(tmp_path / "noswap.docx")
    generate_docx({**MINIMAL_ALLOC, "swap_refs": []}, [], out)
    texts = [p.text for p in Document(out).paragraphs]
    assert any("SWAP: NIL" in t for t in texts)


def test_swap_section_avec_avion(tmp_path):
    """Un avion en swap → section 'SWAP:' + vols copiés depuis autres_rotations (sans titre)."""
    alloc = {
        **MINIMAL_ALLOC,
        "autres_rotations": [
            {
                "ac_reg": "HB-JYA",
                "new_based": False,
                "flights": [
                    _flight("EZS1393", "GVA", "NCE", "0425", "0545", ac_reg="HB-JYA"),
                    _flight("EZS1058", "NCE", "BSL", "0630", "0745", ac_reg="HB-JYA"),
                ],
            },
        ],
        "swap_refs": ["HB-JYA"],
    }
    out = str(tmp_path / "swap.docx")
    generate_docx(alloc, [], out)
    texts = [p.text for p in Document(out).paragraphs]
    assert any("SWAP:" in t and "NIL" not in t for t in texts)
    # L'immat apparaît dans les lignes de vol (avec tirets), pas en titre séparé
    assert any("HB-JYA" in t for t in texts)
    assert any("EZS1393" in t for t in texts)
    assert any("EZS1058" in t for t in texts)


def test_swap_pas_de_nil_si_avion_present(tmp_path):
    alloc = {
        **MINIMAL_ALLOC,
        "autres_rotations": [{
            "ac_reg": "HB-JYA", "new_based": False,
            "flights": [_flight("EZS1393", "GVA", "NCE", ac_reg="HB-JYA")],
        }],
        "swap_refs": ["HB-JYA"],
    }
    out = str(tmp_path / "swap2.docx")
    generate_docx(alloc, [], out)
    texts = [p.text for p in Document(out).paragraphs]
    assert not any("SWAP: NIL" in t for t in texts)


def test_swap_highlights_copies(tmp_path):
    """Les highlights s'appliquent aussi aux lignes copiées dans la section SWAP."""
    alloc = {
        **MINIMAL_ALLOC,
        "autres_rotations": [{
            "ac_reg": "HB-JYA", "new_based": False,
            "flights": [
                _flight("EZS1393", "GVA", "NCE", "0425", "0545", ac_reg="HB-JYA"),
                _flight("EZS1058", "NCE", "BSL", "0630", "0745", ac_reg="HB-JYA"),
            ],
        }],
        "swap_refs": ["HB-JYA"],
    }
    out = str(tmp_path / "swap_hl.docx")
    generate_docx(alloc, [{"flt_no": "EZS1393", "color_index": 0}], out)
    doc2 = Document(out)
    # Le vol EZS1393 doit être surligné (y compris dans la section SWAP)
    highlighted = [
        p for p in doc2.paragraphs
        if "EZS1393" in p.text and p.runs and p.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW
    ]
    assert len(highlighted) == 2  # une fois dans AUTRES ROTATIONS, une fois dans SWAP


def test_ferry_section_si_non_vide(tmp_path):
    alloc = {**MINIMAL_ALLOC, "ferry": [{
        "ac_reg": "G-FERR", "new_based": False,
        "flights": [_flight("EZY9999", "LHR", "BER", ac_reg="G-FERR")],
    }]}
    out = str(tmp_path / "ferry.docx")
    generate_docx(alloc, [], out)
    texts = [p.text for p in Document(out).paragraphs]
    assert any("FERRY" in t for t in texts)


def test_ferry_nil_si_vide(tmp_path):
    """Ferry vide → 'FERRY: NIL' doit apparaître."""
    out = str(tmp_path / "noferry.docx")
    generate_docx({**MINIMAL_ALLOC, "ferry": []}, [], out)
    texts = [p.text for p in Document(out).paragraphs]
    assert any("FERRY: NIL" in t for t in texts)


# ---------------------------------------------------------------------------
# Intégration bout-en-bout sur le PDF réel
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def real_doc():
    from services.pdf_parser import parse_easyjet_allocation
    from services.alloc_builder import build_allocation

    pdf = Path(__file__).resolve().parents[4] / "workflow" / "FlightAllocationReport (1).pdf"
    alloc = build_allocation(parse_easyjet_allocation(str(pdf)))
    highlights = [
        {"flt_no": "EJU1687", "color_index": 0},
        {"flt_no": "EJU5143", "color_index": 3},
    ]
    out = Path(tempfile.mkdtemp()) / "real_alloc.docx"
    generate_docx(alloc, highlights, str(out))
    return Document(str(out))


def test_real_nb_paragraphes(real_doc):
    assert len(real_doc.paragraphs) > 100


def test_real_titre(real_doc):
    assert real_doc.paragraphs[0].text.startswith("Operations EasyJet NCE")
    assert "27 JUN" in real_doc.paragraphs[0].text


def test_real_eju1687_yellow(real_doc):
    p = next(p for p in real_doc.paragraphs if "EJU1687" in p.text)
    assert p.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW


def test_real_eju5143_pink(real_doc):
    p = next(p for p in real_doc.paragraphs if "EJU5143" in p.text)
    assert p.runs[0].font.highlight_color == WD_COLOR_INDEX.PINK


def test_real_swap_section(real_doc):
    texts = [p.text for p in real_doc.paragraphs]
    assert any("SWAP" in t for t in texts)
    # Les immatriculations apparaissent dans les lignes de vol (avec tirets)
    assert any("HB-JYA" in t for t in texts)
    assert any("HB-JZR" in t for t in texts)


def test_real_5_avions_bases(real_doc):
    # Les 5 avions basés NCE apparaissent sans tiret
    base_regs = {"OEICI", "OEIJH", "OEIJR", "OEINN", "OEIVT"}
    texts = {p.text for p in real_doc.paragraphs}
    for reg in base_regs:
        assert any(reg in t for t in texts), f"{reg} introuvable"


# ---------------------------------------------------------------------------
# Night Stop — un avion qui repart ne dort pas sur place
# ---------------------------------------------------------------------------

class TestNightStopReel:
    """Cas relevés sur l'allocation du 30JUL26 : le tableau donnait un night
    stop à un avion reparti dans la journée, et une « first wave » d'après-midi
    à un avion arrivé en cours de journée."""

    @staticmethod
    def _f(flt, dep, arr, std, sta, reg="OE-INI"):
        return {"date": "30.07", "flt_no": flt, "dep": dep, "arr": arr,
                "std": std, "sta": sta, "ac_reg": reg, "ac_type": "320",
                "capacity": 186, "pax": 100, "crew_str": "2+4/0",
                "captain": None, "crew_change_before": False}

    def test_repart_apres_derniere_arrivee_nce(self):
        from services.docx_generator import _night_stop
        vols = [
            self._f("EJU1611", "NCE", "LIL", "05:00", "06:40"),
            self._f("EJU1612", "LIL", "NCE", "07:15", "09:00"),
            self._f("EJU1693", "NCE", "PMO", "09:45", "11:15"),   # part et ne revient pas
        ]
        assert _night_stop(vols) is None

    def test_termine_a_nce(self):
        from services.docx_generator import _night_stop
        vols = [
            self._f("EJU1611", "NCE", "LIL", "05:00", "06:40"),
            self._f("EJU1612", "LIL", "NCE", "07:15", "09:00"),
        ]
        ns = _night_stop(vols)
        assert ns is not None and ns["flt_no"] == "EJU1612"

    def test_sans_night_stop_pas_de_first_wave(self, tmp_path):
        """Un avion arrivé en cours de journée n'a pas de first wave : sa
        colonne doit rester vide plutôt que d'afficher un départ d'après-midi."""
        from docx import Document
        from services.docx_generator import generate_docx
        entry = {"ac_reg": "OE-ICM", "new_based": True, "flights": [
            self._f("EJU1694", "PMO", "NCE", "12:15", "13:50", "OE-ICM"),
            self._f("EJU1697", "NCE", "VCE", "14:50", "16:00", "OE-ICM"),
        ]}
        alloc = {"date": "30 JUL", "based": [entry], "autres_rotations": [],
                 "swap_refs": [], "ferry": []}
        out = str(tmp_path / "a.docx")
        generate_docx(alloc, [], out, previous_alloc_flights=[])  # pas de veille
        row = [c.text for c in Document(out).tables[0].rows[1].cells]
        assert row[0] == "" and row[1] == "", "night stop inexistant"
        assert row[2] == "" and row[3] == "", "aucune first wave sans night stop"
        assert row[5] == "OEICM"
