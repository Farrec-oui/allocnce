"""Integration tests — real PDF files from /Users/thomaszahnd/Desktop/alloc/workflow/."""
import os
import re
import tempfile
from pathlib import Path

import pytest
from docx import Document

from services.alloc_builder import build_allocation
from services.comparator import (
    compare_allocs,
    compare_with_feuille_journee,
    parse_feuille_journee,
)
from services.docx_generator import generate_docx
from services.pdf_parser import parse_easyjet_allocation

WORKFLOW = Path("/Users/thomaszahnd/Desktop/alloc/workflow")
PDF_ALLOC_1 = WORKFLOW / "FlightAllocationReport.pdf"
PDF_ALLOC_2 = WORKFLOW / "FlightAllocationReport (1).pdf"
PDF_FJ = WORKFLOW / "Feuille de journee EASYJET MZS V19 - NEW_fr-fr.pdf"

FLT_RE = re.compile(r"^[A-Z]{2,3}\d{4}$")


# ---------------------------------------------------------------------------
# Module-level fixtures (parsed once, shared across all tests)
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def flights_1():
    return parse_easyjet_allocation(str(PDF_ALLOC_1))


@pytest.fixture(scope="module")
def alloc_1(flights_1):
    return build_allocation(flights_1)


@pytest.fixture(scope="module")
def flights_2():
    return parse_easyjet_allocation(str(PDF_ALLOC_2))


@pytest.fixture(scope="module")
def flights_fj():
    return parse_feuille_journee(str(PDF_FJ))


# ---------------------------------------------------------------------------
# TEST 1 — parse_easyjet_allocation
# ---------------------------------------------------------------------------

class TestParsePdf:
    def test_au_moins_60_vols(self, flights_1):
        assert len(flights_1) >= 60, f"Attendu ≥60 vols, obtenu {len(flights_1)}"

    def test_eju1687_dep(self, flights_1):
        eju1687 = next(f for f in flights_1 if f["flt_no"] == "EJU1687")
        assert eju1687["dep"] == "NCE"

    def test_eju1687_arr(self, flights_1):
        eju1687 = next(f for f in flights_1 if f["flt_no"] == "EJU1687")
        assert eju1687["arr"] == "OLB"

    def test_eju1687_std(self, flights_1):
        eju1687 = next(f for f in flights_1 if f["flt_no"] == "EJU1687")
        assert eju1687["std"] == "04:00"

    def test_eju1687_ac_reg(self, flights_1):
        eju1687 = next(f for f in flights_1 if f["flt_no"] == "EJU1687")
        assert eju1687["ac_reg"] == "OE-IJR"

    def test_tri_std_croissant(self, flights_1):
        stds = [f["std"] for f in flights_1]
        assert stds == sorted(stds), "Les vols ne sont pas triés par STD croissant"


# ---------------------------------------------------------------------------
# TEST 2 — build_allocation
# ---------------------------------------------------------------------------

class TestBuildAllocation:
    def test_oeici_dans_based(self, alloc_1):
        based_regs = [e["ac_reg"] for e in alloc_1["based"]]
        assert "OE-ICI" in based_regs

    def test_based_trie_alphabetiquement(self, alloc_1):
        based_regs = [e["ac_reg"] for e in alloc_1["based"]]
        assert based_regs == sorted(based_regs), "Les avions 'based' ne sont pas triés"

    def test_oeici_rotations_completes(self, alloc_1):
        oeici = next(e for e in alloc_1["based"] if e["ac_reg"] == "OE-ICI")
        flt_nos = [f["flt_no"] for f in oeici["flights"]]
        # OE-ICI fait NCE→HRG, HRG→NCE, NCE→FNC, FNC→NCE
        assert "EJU1663" in flt_nos
        assert "EJU1664" in flt_nos
        assert "EJU1745" in flt_nos
        assert "EJU1746" in flt_nos

    def test_sections_presentes(self, alloc_1):
        assert isinstance(alloc_1["based"], list)
        assert isinstance(alloc_1["autres_rotations"], list)
        assert isinstance(alloc_1["swap_refs"], list)
        assert isinstance(alloc_1["ferry"], list)

    def test_date_extraite(self, alloc_1):
        assert alloc_1["date"] == "27 JUN"


# ---------------------------------------------------------------------------
# Module-level fixtures for tests 3–5 (avoid PytestRemovedIn10Warning)
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def docx_path(alloc_1):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    generate_docx(alloc_1, [], path)
    yield path
    os.unlink(path)


@pytest.fixture(scope="module")
def highlights_compare(flights_1, flights_2):
    return compare_allocs(flights_1, flights_2, [], 1)["highlights"]


@pytest.fixture(scope="module")
def prealloc_docx(flights_1, alloc_1, flights_fj):
    highlights = compare_with_feuille_journee(flights_1, flights_fj)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    generate_docx(alloc_1, highlights, path)
    yield path, highlights
    os.unlink(path)


# ---------------------------------------------------------------------------
# TEST 3 — generate_docx
# ---------------------------------------------------------------------------

class TestGenerateDocx:
    def test_fichier_existe(self, docx_path):
        assert os.path.exists(docx_path)

    def test_taille_sup_5ko(self, docx_path):
        size = os.path.getsize(docx_path)
        assert size > 5_000, f"DOCX trop petit : {size} octets"

    def test_oeici_present_comme_paragraphe(self, docx_path):
        doc = Document(docx_path)
        texts = [p.text.strip() for p in doc.paragraphs]
        assert "OEICI" in texts, "Paragraphe 'OEICI' introuvable dans le DOCX"

    def test_entete_present(self, docx_path):
        doc = Document(docx_path)
        texts = [p.text.strip() for p in doc.paragraphs]
        assert any("27 JUN" in t for t in texts), "Date '27 JUN' absente du DOCX"

    def test_vol_eju1687_present(self, docx_path):
        doc = Document(docx_path)
        texts = [p.text for p in doc.paragraphs]
        assert any("EJU1687" in t for t in texts)


# ---------------------------------------------------------------------------
# TEST 4 — compare_allocs entre les deux PDFs réels
# ---------------------------------------------------------------------------

class TestCompareAllocs:
    def test_highlights_non_vides(self, highlights_compare):
        assert len(highlights_compare) > 0, "Aucun highlight détecté entre les deux PDFs"

    def test_tous_flt_no_valides(self, highlights_compare):
        invalides = [h["flt_no"] for h in highlights_compare if not FLT_RE.match(h["flt_no"])]
        assert invalides == [], f"flt_no invalides : {invalides}"

    def test_color_index_correct(self, highlights_compare):
        assert all(h["color_index"] == 1 for h in highlights_compare)

    def test_ezs1055_ou_1056_detecte(self, highlights_compare):
        flt_nos = {h["flt_no"] for h in highlights_compare}
        assert flt_nos & {"EZS1055", "EZS1056"}, (
            f"EZS1055/EZS1056 attendus dans highlights, obtenu : {flt_nos}"
        )


# ---------------------------------------------------------------------------
# TEST 5 — workflow complet PreAlloc
# ---------------------------------------------------------------------------

class TestWorkflowPrealloc:
    def test_docx_taille_raisonnable(self, prealloc_docx):
        path, _ = prealloc_docx
        assert os.path.getsize(path) > 5_000

    def test_oeici_present(self, prealloc_docx):
        path, _ = prealloc_docx
        doc = Document(path)
        texts = [p.text.strip() for p in doc.paragraphs]
        assert "OEICI" in texts

    def test_sections_presentes(self, prealloc_docx):
        path, _ = prealloc_docx
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Night Stop" in all_text
        assert "Rotations" in all_text

    def test_highlights_sont_des_flt_no_valides(self, prealloc_docx):
        _, highlights = prealloc_docx
        invalides = [h["flt_no"] for h in highlights if not FLT_RE.match(h["flt_no"])]
        assert invalides == [], f"flt_no invalides : {invalides}"

    def test_meme_avions_based_que_sans_fj(self, prealloc_docx, alloc_1):
        path, _ = prealloc_docx
        doc = Document(path)
        texts = [p.text.strip() for p in doc.paragraphs]
        based_regs = [e["ac_reg"] for e in alloc_1["based"]]
        for reg in based_regs:
            reg_tag = reg.replace("-", "")
            assert reg_tag in texts, f"Avion basé '{reg_tag}' absent du DOCX prealloc"
