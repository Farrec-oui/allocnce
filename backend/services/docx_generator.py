import os
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLOR_MAP: dict[int, WD_COLOR_INDEX] = {
    0: WD_COLOR_INDEX.YELLOW,
    1: WD_COLOR_INDEX.BRIGHT_GREEN,
    2: WD_COLOR_INDEX.TURQUOISE,
    3: WD_COLOR_INDEX.PINK,
}

SEPARATOR  = "*" * 58
_HDR_GRAY  = "D9D9D9"


# ---------------------------------------------------------------------------
# Helpers texte bas niveau
# ---------------------------------------------------------------------------

def _set_run(run, size: int = 9, hl: WD_COLOR_INDEX | None = None) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if hl is not None:
        run.font.highlight_color = hl


def _para(doc: Document, text: str = "", size: int = 9,
          hl: WD_COLOR_INDEX | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if text:
        _set_run(p.add_run(text), size=size, hl=hl)
    return p


def _para_flight(
    doc: Document,
    f: dict,
    hl_info: tuple[WD_COLOR_INDEX, bool] | None,
) -> None:
    """Write one flight line paragraph with full or partial highlighting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if hl_info is None:
        _set_run(p.add_run(_flight_line(f)))
    elif not hl_info[1]:
        # Full highlight — one run, entire line
        _set_run(p.add_run(_flight_line(f)), hl=hl_info[0])
    else:
        # Partial highlight — only the ac_reg token is highlighted
        prefix, immat, suffix = _flight_line_split(f)
        _set_run(p.add_run(prefix))
        _set_run(p.add_run(immat), hl=hl_info[0])
        _set_run(p.add_run(suffix))


def _para_ac_label(doc: Document, text: str) -> None:
    """Titre d'immatriculation : gras, rouge, souligné."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.font.name      = "Calibri"
    run.font.size      = Pt(9)
    run.bold           = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.underline      = True


def _para_crew_change(doc: Document) -> None:
    """Ligne 'Crew change' : rouge, souligné (pas gras)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run("Crew change")
    run.font.name      = "Calibri"
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.underline      = True


def _para_legend(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(
        "En surbrillance les modifications apportées aux précédents envois "
    )
    _set_run(r1)
    r2 = p.add_run(" ")
    _set_run(r2, hl=WD_COLOR_INDEX.YELLOW)


# ---------------------------------------------------------------------------
# Formatage des données de vol
# ---------------------------------------------------------------------------

def _time(t: str) -> str:
    return t.replace(":", "")


def _ac_label(ac_reg: str, new_based: bool) -> str:
    label = ac_reg.replace("-", "")
    return f"{label}  New based" if new_based else label


def _flight_line(f: dict) -> str:
    # Raw PDF flights (e.g. cancelled) may lack capacity/crew_str; fall back to pax/crew
    capacity = f.get("capacity", "")
    crew_str = f.get("crew_str") or f.get("crew", "")
    parts = [
        f["date"], f["flt_no"], f["dep"], f["arr"],
        _time(f["std"]), _time(f["sta"]),
        f["ac_reg"], f["ac_type"],
        str(capacity), str(f["pax"]),
        crew_str,
    ]
    line = " ".join(parts)
    if f.get("captain"):
        line += " " + f["captain"][:8]
    return line


def _flight_line_split(f: dict) -> tuple[str, str, str]:
    """Return (prefix, ac_reg, suffix) for partial immat-only highlighting.

    prefix  = 'DD.MM FLT DEP ARR STD STA '
    ac_reg  = 'OE-ICI'
    suffix  = ' TYPE CAP PAX CREW [captain]'
    """
    prefix = " ".join([
        f["date"], f["flt_no"], f["dep"], f["arr"],
        _time(f["std"]), _time(f["sta"]),
    ]) + " "
    immat = f["ac_reg"]
    capacity = f.get("capacity", "")
    crew_str = f.get("crew_str") or f.get("crew", "")
    suffix = " " + " ".join([
        f["ac_type"], str(capacity), str(f["pax"]), crew_str,
    ])
    if f.get("captain"):
        suffix += " " + f["captain"][:8]
    return prefix, immat, suffix


# ---------------------------------------------------------------------------
# Écriture des sections texte (vols)
# ---------------------------------------------------------------------------

def _write_section(
    doc: Document,
    entries: list[dict],
    hl_map: dict[str, tuple[WD_COLOR_INDEX, bool]],
) -> None:
    """Section BASÉS : titre ac_reg + ligne vide + vols + ligne vide."""
    for entry in entries:
        _para_ac_label(doc, _ac_label(entry["ac_reg"], entry.get("new_based", False)))
        _para(doc)
        for f in entry["flights"]:
            if f.get("crew_change_before"):
                _para_crew_change(doc)
            _para_flight(doc, f, hl_map.get(f["flt_no"]))
        _para(doc)


def _write_autres_rotations(
    doc: Document,
    entries: list[dict],
    hl_map: dict[str, tuple[WD_COLOR_INDEX, bool]],
) -> None:
    """AUTRES ROTATIONS / SWAP / FERRY : vols SANS titre d'immatriculation."""
    for entry in entries:
        for f in entry["flights"]:
            if f.get("crew_change_before"):
                _para_crew_change(doc)
            _para_flight(doc, f, hl_map.get(f["flt_no"]))
        _para(doc)


# ---------------------------------------------------------------------------
# Helpers tableaux Word
# ---------------------------------------------------------------------------

def _set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    centered: bool = False,
) -> None:
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0] if p.runs else p.add_run()
    run.text       = text
    run.font.name  = "Calibri"
    run.font.size  = Pt(9)
    run.font.bold  = bold


def _set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_table_borders(table) -> None:
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")       # 0.5pt = 4 eighths
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Helpers données Night Stop / First Wave
# ---------------------------------------------------------------------------

def _find_first_nce_dep(entry: dict) -> dict | None:
    for f in entry["flights"]:
        if f["dep"] == "NCE":
            return f
    return None


def _night_stop(flights: list[dict]) -> dict | None:
    """Vol par lequel l'avion termine sa journée à NCE, s'il y reste.

    Prendre la dernière arrivée à NCE ne suffit pas : un avion peut y arriver
    puis repartir sans revenir — il ne passe alors pas la nuit sur place. Seul
    compte le dernier mouvement de la journée : si c'est une arrivée à NCE,
    l'avion y dort ; si c'est un départ, il est ailleurs.
    """
    if not flights:
        return None
    last = sorted(flights, key=lambda f: f["std"])[-1]
    return last if last["arr"] == "NCE" else None


def _night_stops_by_reg(flights: list[dict]) -> dict[str, dict]:
    """{ac_reg: vol de night stop} pour une liste de vols brute (une journée)."""
    by_reg: dict[str, list[dict]] = defaultdict(list)
    for f in flights:
        by_reg[f["ac_reg"]].append(f)
    out: dict[str, dict] = {}
    for ac_reg, ac_flights in by_reg.items():
        ns = _night_stop(ac_flights)
        if ns is not None:
            out[ac_reg] = ns
    return out


# ---------------------------------------------------------------------------
# Tableau 0 — Night Stop J / First Wave J+1
# ---------------------------------------------------------------------------

def _add_table0(
    doc: Document,
    alloc_data: dict,
    previous_alloc_flights: list[dict] | None,
) -> None:
    headers  = ["Night Stop", "STA", "First Wave", "STD", "Stand", "Reg"]
    widths   = [Cm(5), Cm(1.5), Cm(5), Cm(1.5), Cm(1.5), Cm(1.5)]
    centered = {1, 3, 4, 5}

    based = alloc_data.get("based", [])

    # Avions ayant réellement passé la nuit précédente à NCE
    prev_ns = _night_stops_by_reg(previous_alloc_flights or [])

    n_rows = 1 + max(len(based), 1)
    table  = doc.add_table(rows=n_rows, cols=6)
    _set_table_borders(table)

    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_text(cell, h, bold=True, centered=(j in centered))
        _set_cell_bg(cell, _HDR_GRAY)

    # Data rows
    for i, entry in enumerate(based):
        ac_reg = entry["ac_reg"]
        ns_f   = prev_ns.get(ac_reg)
        # Sans night stop la veille, il n'y a pas de « first wave » : le premier
        # départ de l'avion n'est alors pas la reprise du matin (un avion arrivé
        # en cours de journée repartirait par exemple à 14h50).
        fw_f   = _find_first_nce_dep(entry) if ns_f else None

        ns_text  = f"{ns_f['date']} {ns_f['flt_no']} {ns_f['dep']} {ns_f['arr']}" if ns_f else ""
        sta_text = _time(ns_f["sta"]) if ns_f else ""
        fw_text  = f"{fw_f['date']} {fw_f['flt_no']} {fw_f['dep']} {fw_f['arr']}" if fw_f else ""
        std_text = _time(fw_f["std"]) if fw_f else ""
        reg_text = ac_reg.replace("-", "")

        for j, val in enumerate([ns_text, sta_text, fw_text, std_text, "", reg_text]):
            _set_cell_text(table.cell(i + 1, j), val, centered=(j in centered))


# ---------------------------------------------------------------------------
# Tableau 1 — Départs Remote stand
# ---------------------------------------------------------------------------

def _add_table1(doc: Document) -> None:
    headers = [
        "VOL DEP", "STD", "Reg.", "PKG",
        "Porte embarquement (Confirmée avec ACA)",
        "Nombre BUS", "Heure d'embarquement (H-50min)",
    ]
    widths   = [Cm(2.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(4.5), Cm(2.0), Cm(3.0)]
    centered = {1, 2, 3, 5, 6}

    table = doc.add_table(rows=2, cols=7)
    _set_table_borders(table)

    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_text(cell, h, bold=True, centered=(j in centered))
        _set_cell_bg(cell, _HDR_GRAY)
    # Row 1 left empty for manual input


# ---------------------------------------------------------------------------
# Tableau 2 — First Wave J+2 (demain soir)
# ---------------------------------------------------------------------------

def _add_table2(doc: Document, alloc_data: dict) -> None:
    headers  = ["Night Stop", "STA", "First Wave", "STD", "Reg."]
    widths   = [Cm(5), Cm(1.5), Cm(3), Cm(1.5), Cm(1.5)]
    centered = {1, 2, 3, 4}

    based  = alloc_data.get("based", [])
    n_rows = 1 + max(len(based), 1)
    table  = doc.add_table(rows=n_rows, cols=5)
    _set_table_borders(table)

    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_text(cell, h, bold=True, centered=(j in centered))
        _set_cell_bg(cell, _HDR_GRAY)

    for i, entry in enumerate(based):
        ac_reg = entry["ac_reg"]
        ns_f   = _night_stop(entry["flights"])

        ns_text  = f"{ns_f['date']} {ns_f['flt_no']} {ns_f['dep']} {ns_f['arr']}" if ns_f else ""
        sta_text = _time(ns_f["sta"]) if ns_f else ""
        fw_text  = "TBA" if ns_f else ""
        reg_text = ac_reg.replace("-", "")

        for j, val in enumerate([ns_text, sta_text, fw_text, "", reg_text]):
            _set_cell_text(table.cell(i + 1, j), val, centered=(j in centered))


# ---------------------------------------------------------------------------
# Point d'entrée public
# ---------------------------------------------------------------------------

def generate_docx(
    alloc_data: dict,
    highlights: list[dict],
    output_path: str,
    previous_alloc_flights: list[dict] | None = None,
    cancelled: list[dict] | None = None,
) -> str:
    """Génère le document .docx d'allocation.

    Args:
        alloc_data:              sortie de build_allocation()
        highlights:              [{"flt_no": "EJU1604", "color_index": 0}, ...]
        output_path:             chemin de destination
        previous_alloc_flights:  vols bruts de l'alloc du jour J (la veille),
                                 utilisés pour remplir la colonne Night Stop
                                 du Tableau 0. None → colonnes Night Stop vides.
    """
    hl_map: dict[str, tuple[WD_COLOR_INDEX, bool]] = {
        h["flt_no"]: (COLOR_MAP[h["color_index"] % 4], h.get("partial", False))
        for h in (highlights or [])
    }

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Cm(1)
    sec.bottom_margin = Cm(1)
    sec.left_margin   = Cm(2)
    sec.right_margin  = Cm(2)

    # 1. Titre + légende
    date = alloc_data.get("date", "")
    _para(doc, f"Operations EasyJet NCE  {date} Toutes les heures sont en Z", size=11)
    _para_legend(doc)

    # 2. Ligne vide
    _para(doc)

    # 3. Tableau Night Stop J / First Wave J+1
    _para(doc, "Night Stop J / First Wave J+1 ")
    _add_table0(doc, alloc_data, previous_alloc_flights)
    _para(doc)
    _para(doc)

    # 5. Tableau Départs Remote stand
    _para(doc, "Départs Remote stand :")
    _add_table1(doc)
    _para(doc)

    # 7. Tableau First Wave J+2
    _para(doc, "First Wave J+2 (demain soir)")
    _add_table2(doc, alloc_data)
    _para(doc)
    _para(doc)

    # 9. Section rotations
    _para(doc, "Rotations de la journée + Relève équipage de nos appareils basés ")
    _para(doc)

    # 11. Basés
    _write_section(doc, alloc_data.get("based", []), hl_map)

    # 12. Séparateur
    _para(doc, SEPARATOR)

    # 13-14. Autres rotations
    _para(doc, "AUTRES ROTATIONS: ")
    _para(doc)
    _write_autres_rotations(doc, alloc_data.get("autres_rotations", []), hl_map)

    # 15-16. SWAP (séparateur seulement si non NIL)
    swap_refs: list[str] = alloc_data.get("swap_refs", [])
    if swap_refs:
        _para(doc, SEPARATOR)
        _para(doc, "SWAP:")
        swap_set     = set(swap_refs)
        swap_entries = sorted(
            (e for e in alloc_data.get("autres_rotations", []) if e["ac_reg"] in swap_set),
            key=lambda e: e["ac_reg"],
        )
        _write_autres_rotations(doc, swap_entries, hl_map)
    else:
        _para(doc, "SWAP: NIL")

    # 17. FERRY (toujours présent)
    if alloc_data.get("ferry"):
        _para(doc, "FERRY:")
        _write_autres_rotations(doc, alloc_data["ferry"], hl_map)
    else:
        _para(doc, "FERRY: NIL")

    # 18. CNL (si non vide)
    if cancelled:
        _para(doc, SEPARATOR)
        _para(doc, "CNL :")
        _para(doc)
        for f in sorted(cancelled, key=lambda x: x.get("std", "")):
            _para(doc, _flight_line(f))

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path
